from docx import Document
import docx
import itertools

# Sentence similarity function


def similarity_sentence_compare(before_doc, after_doc, user_name, save_list):
    # 각 Sentence의 index를 담는 list 생성
    same_sentence_line = []
    new_sentence_line = []
    moved_sentence_line = []
    save_list.clear()
    before = list(map(lambda x: x.text.split("."), before_doc.paragraphs[:]))
    before = list(itertools.chain(*before))
    while '' in before:
        before.remove('')
    while ' ' in a:
        before.remove(' ')

    after = list(map(lambda x: x.text.split("."), after_doc.paragraphs[:]))
    after = list(itertools.chain(*after))
    while '' in before:
        after.remove('')
    while ' ' in a:
        after.remove(' ')

    # Sentence간 비교 (after file <-> before file)
    for x, after_sentence in enumerate(after):
        # 첫번째 for문(After File)의 각각 Sentence에 대해서,
        # 두번째 for문(Before File) 전체의 Sentence를 순회하며 비교
        for y, before_sentence in enumerate(before):
            # Before과 After Sentence가 동일한 경우 (줄바꿈은 '' 빈 문자열로, 모든 줄바꿈이 == 하므로 elif문으로 뺌)
            if before_sentence == after_sentence:
                # Index 역시 동일하면 same list에 추가
                if x == y:
                    same_sentence_line.append(x)
                    break
                # Index가 동일하지 않으면 moved list에 추가 (before 문서의 기존 Sentence가 위치만 이동)
                else:
                    moved_sentence_line.append(x)
                    break

        # 전체 순회가 끝났으나, 동일한 Sentence가 발견되지 않은 경우 => 새롭게 추가된/수정된 Sentence
        if x not in same_sentence_line and x not in moved_sentence_line:
            new_sentence_line.append(x)

    # 전체에서 new_sentence_line의 비율 계산
    ratio_new_sentence = round(
        len(new_sentence_line) /
        len(same_sentence_line + moved_sentence_line + new_sentence_line) *
        100, 2)

    save_list.append(user_name)
    save_list.append(same_sentence_line)
    save_list.append(moved_sentence_line)
    save_list.append(new_sentence_line)
    save_list.append(ratio_new_sentence)

# 비교할 before, after 문서 파일, 최종적으로 함께 list에 넣어줄 user_name, 결과를 받을 빈 list를 넣어서 사용


def similarity_compare(before_doc, after_doc, user_name, save_list):
    # 각 paragraph의 index를 담는 list 생성
    same_paragraph_line = []
    new_paragraph_line = []
    moved_parahraph_line = []
    save_list.clear()
    # paragraph간 비교 (after file <-> before file)
    for x, paragraph_after in enumerate(after_doc.paragraphs):
        # 첫번째 for문(After File)의 각각 Paragraph에 대해서,
        # 두번째 for문(Before File) 전체의 Paragraph를 순회하며 비교
        for y, paragraph_before in enumerate(before_doc.paragraphs):
            # Before과 After Paragraph가 동일한 경우 (줄바꿈은 '' 빈 문자열로, 모든 줄바꿈이 == 하므로 elif문으로 뺌)
            if paragraph_before.text == paragraph_after.text and paragraph_after.text != '':
                # Index 역시 동일하면 same list에 추가
                if x == y:
                    same_paragraph_line.append(x)
                    break
                # Index가 동일하지 않으면 moved list에 추가 (before 문서의 기존 paragraph가 위치만 이동)
                else:
                    moved_parahraph_line.append(x)
                    break
            #
            elif paragraph_after.text == '':
                if paragraph_before.text == '':
                    same_paragraph_line.append(x)
                    break

        # 전체 순회가 끝났으나, 동일한 Paragraph가 발견되지 않은 경우 => 새롭게 추가된/수정된 Paragraph
        if x not in same_paragraph_line and x not in moved_parahraph_line:
            new_paragraph_line.append(x)

    # 전체에서 new_paragraph_line의 비율 계산 (전체 문서에서 새롭게 추가된 내용의 비율)
    ratio_new_para = round(
        len(new_paragraph_line) /
        len(same_paragraph_line + moved_parahraph_line + new_paragraph_line) *
        100, 2)

    save_list.append(user_name)
    save_list.append(same_paragraph_line)
    save_list.append(moved_parahraph_line)
    save_list.append(new_paragraph_line)
    save_list.append(ratio_new_para)


def show(document, index):  # Target document와 조회할 index
    return document.paragraphs[index].text
