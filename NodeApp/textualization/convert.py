from docx import Document
import docx


def convert(target_document):
    # FileName Extension
    # fn_extension = target_document.split(".")[-1]
    fn_extension = "docx"  # 함수 내에서 처리하면 오류 발생하여 일단 all true로 처리함
    if fn_extension == "docx":
        return docx.Document(target_document)


def show_paragraph(converted_document, index):
    docu_type_str = str(type(converted_document))
    if "docx.document.Document" in docu_type_str:
        if index == "all":
            for x, paragraph in enumerate(converted_document.paragraphs):
                print(str(x) + " : " + paragraph.text)
            return None
        elif type(index) == int:
            return converted_document.paragraphs[index].text
    else:
        raise Exception("Unsupported file extension. Only use docx or ...")


def get_str(converted_document):
    docu_type_str = str(type(converted_document))
    if "docx.document.Document" in docu_type_str:
        document_string = ""
        for paragraph in converted_document.paragraphs:
            document_string = document_string + paragraph.text.rstrip()
        return document_string
    else:
        raise Exception("Unsupported file extension. Only use docx or ...")


# docx file의 txt를 DB에 삽입하거나 별도 이용시 에러 발생시키는 이스케이프 코드 제거한 후, str 반환
def docx_read(file=None):
    txt_val = ''
    if file:
        for paragraph in file.paragraphs:
            if paragraph.text == '':  # 여러줄의 공백 제거를 위한 조건문
                pass
            else:
                txt_val = txt_val + paragraph.text
    else:
        return 0

    escape_code = ["\n", "\t", "\'", "“", '”', '\/', '/', '\"', '\r', '\b']
    for escape in escape_code:
        txt_val = txt_val.replace(escape, "")

    # 전체 문자열 안에 \n이 포함되어 있으므로, print 함수로 씌워서 함수 실행하면 줄바꿈 되어 출력
    return str(txt_val)